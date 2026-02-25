from django.http import HttpResponse
from django.utils import timezone
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

from interventions.models import InterventionCase
from alerts.models import Alert
from students.models import Student, Classroom
from users.models import User
from django.db.models import Count, Avg, F
from datetime import timedelta


class ReportGenerator:
    """Professional report generator for Somali Early Warning System"""
    
    def __init__(self):
        self.title = "Somali Early Warning System"
        self.subtitle = "School Support & Intervention Report"
        
    def generate_pdf_report(self, report_type):
        """Generate professional PDF report"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=14,
            textColor=colors.HexColor('#64748b'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        # Header
        story.append(Paragraph(self.title, title_style))
        story.append(Paragraph(self.subtitle, subtitle_style))
        story.append(Paragraph(f"Generated: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}", 
                               ParagraphStyle('Date', parent=styles['Normal'], fontSize=10, 
                                            textColor=colors.grey, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.3*inch))
        
        # Report content based on type
        if report_type == 'cases':
            story.extend(self._generate_cases_content_pdf(styles))
        elif report_type == 'risk':
            story.extend(self._generate_risk_content_pdf(styles))
        elif report_type == 'performance':
            story.extend(self._generate_performance_content_pdf(styles))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _generate_cases_content_pdf(self, styles):
        """Generate intervention cases content for PDF"""
        content = []
        
        # Section header
        section_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        content.append(Paragraph("Intervention Cases Report", section_style))
        content.append(Spacer(1, 0.2*inch))
        
        # Get data
        cases = InterventionCase.objects.select_related('student', 'assigned_to').all()
        
        # Summary statistics
        total_cases = cases.count()
        open_cases = cases.exclude(status='closed').count()
        closed_cases = cases.filter(status='closed').count()
        
        summary_data = [
            ['Metric', 'Value'],
            ['Total Cases', str(total_cases)],
            ['Open Cases', str(open_cases)],
            ['Closed Cases', str(closed_cases)],
            ['Success Rate', f'{(closed_cases/total_cases*100):.1f}%' if total_cases > 0 else '0%']
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        
        content.append(summary_table)
        content.append(Spacer(1, 0.3*inch))
        
        # Detailed cases table
        content.append(Paragraph("Detailed Case List", section_style))
        content.append(Spacer(1, 0.1*inch))
        
        case_data = [['Case ID', 'Student', 'Form Master', 'Status', 'Days Open']]
        
        for case in cases[:20]:  # Limit to 20 for PDF
            days_open = (timezone.now().date() - case.created_at.date()).days
            case_data.append([
                str(case.case_id),
                case.student.full_name[:20],
                (case.assigned_to.name[:15] if case.assigned_to else 'Unassigned'),
                case.status,
                str(days_open)
            ])
        
        case_table = Table(case_data, colWidths=[0.8*inch, 2*inch, 1.5*inch, 1.2*inch, 0.8*inch])
        case_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')])
        ]))
        
        content.append(case_table)
        
        return content
    
    def _generate_risk_content_pdf(self, styles):
        """Generate risk summary content for PDF"""
        content = []
        
        section_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#dc2626'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        content.append(Paragraph("Risk Summary by Classroom", section_style))
        content.append(Spacer(1, 0.2*inch))
        
        classrooms = Classroom.objects.filter(is_active=True).select_related('form_master')
        
        risk_data = [['Classroom', 'Form Master', 'Students', 'High Risk', 'Alerts', 'Risk %']]
        
        for classroom in classrooms:
            students = Student.objects.filter(
                enrollments__classroom=classroom,
                enrollments__is_active=True
            )
            
            total_students = students.count()
            high_risk = students.filter(risk_scores__risk_level__in=['high', 'critical']).distinct().count()
            active_alerts = Alert.objects.filter(student__in=students, status='active').count()
            risk_pct = (high_risk / total_students * 100) if total_students > 0 else 0
            
            risk_data.append([
                classroom.name[:15],
                (classroom.form_master.name[:15] if classroom.form_master else 'N/A'),
                str(total_students),
                str(high_risk),
                str(active_alerts),
                f'{risk_pct:.1f}%'
            ])
        
        risk_table = Table(risk_data, colWidths=[1.5*inch, 1.5*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch])
        risk_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dc2626')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#fef2f2')])
        ]))
        
        content.append(risk_table)
        
        return content
    
    def _generate_performance_content_pdf(self, styles):
        """Generate performance metrics content for PDF"""
        content = []
        
        section_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#16a34a'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        content.append(Paragraph("Form Master Performance Metrics", section_style))
        content.append(Spacer(1, 0.2*inch))
        
        form_masters = User.objects.filter(role='form_master')
        
        perf_data = [['Form Master', 'Classroom', 'Active Cases', 'Avg Days', 'On-Time %', 'Rating']]
        
        for fm in form_masters:
            classroom = Classroom.objects.filter(form_master=fm, is_active=True).first()
            cases = InterventionCase.objects.filter(assigned_to=fm)
            active_cases = cases.exclude(status='closed').count()
            
            closed_cases = cases.filter(status='closed')
            avg_resolution = closed_cases.aggregate(avg_days=Avg(F('updated_at') - F('created_at')))['avg_days']
            avg_days = avg_resolution.days if avg_resolution else 0
            
            on_time = closed_cases.filter(updated_at__lte=F('created_at') + timedelta(days=14)).count()
            on_time_pct = (on_time / closed_cases.count() * 100) if closed_cases.count() > 0 else 0
            
            rating = 'Excellent' if on_time_pct >= 80 else 'Good' if on_time_pct >= 60 else 'Fair'
            
            perf_data.append([
                fm.name[:20],
                (classroom.name[:15] if classroom else 'N/A'),
                str(active_cases),
                str(round(avg_days, 1)),
                f'{on_time_pct:.1f}%',
                rating
            ])
        
        perf_table = Table(perf_data, colWidths=[1.8*inch, 1.5*inch, 1*inch, 0.8*inch, 0.9*inch, 1*inch])
        perf_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#16a34a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0fdf4')])
        ]))
        
        content.append(perf_table)
        
        return content
    
    def generate_docx_report(self, report_type):
        """Generate professional DOCX report"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Title
        title = doc.add_heading(self.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(30, 64, 175)
        
        # Subtitle
        subtitle = doc.add_paragraph(self.subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 116, 139)
        
        # Date
        date_para = doc.add_paragraph(f"Generated: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Spacer
        
        # Report content
        if report_type == 'cases':
            self._generate_cases_content_docx(doc)
        elif report_type == 'risk':
            self._generate_risk_content_docx(doc)
        elif report_type == 'performance':
            self._generate_performance_content_docx(doc)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _generate_cases_content_docx(self, doc):
        """Generate cases content for DOCX"""
        doc.add_heading('Intervention Cases Report', 1)
        
        cases = InterventionCase.objects.select_related('student', 'assigned_to').all()
        
        # Summary
        doc.add_heading('Summary Statistics', 2)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Value'
        table.rows[1].cells[0].text = 'Total Cases'
        table.rows[1].cells[1].text = str(cases.count())
        table.rows[2].cells[0].text = 'Open Cases'
        table.rows[2].cells[1].text = str(cases.exclude(status='closed').count())
        table.rows[3].cells[0].text = 'Closed Cases'
        closed = cases.filter(status='closed').count()
        table.rows[3].cells[1].text = str(closed)
        table.rows[4].cells[0].text = 'Success Rate'
        table.rows[4].cells[1].text = f'{(closed/cases.count()*100):.1f}%' if cases.count() > 0 else '0%'
        
        doc.add_paragraph()
        
        # Detailed list
        doc.add_heading('Detailed Case List', 2)
        case_table = doc.add_table(rows=1, cols=5)
        case_table.style = 'Light Grid Accent 1'
        
        hdr_cells = case_table.rows[0].cells
        hdr_cells[0].text = 'Case ID'
        hdr_cells[1].text = 'Student'
        hdr_cells[2].text = 'Form Master'
        hdr_cells[3].text = 'Status'
        hdr_cells[4].text = 'Days Open'
        
        for case in cases[:30]:
            row_cells = case_table.add_row().cells
            row_cells[0].text = str(case.case_id)
            row_cells[1].text = case.student.full_name
            row_cells[2].text = case.assigned_to.name if case.assigned_to else 'Unassigned'
            row_cells[3].text = case.status
            row_cells[4].text = str((timezone.now().date() - case.created_at.date()).days)
    
    def _generate_risk_content_docx(self, doc):
        """Generate risk content for DOCX"""
        doc.add_heading('Risk Summary by Classroom', 1)
        
        classrooms = Classroom.objects.filter(is_active=True).select_related('form_master')
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Classroom'
        hdr_cells[1].text = 'Form Master'
        hdr_cells[2].text = 'Students'
        hdr_cells[3].text = 'High Risk'
        hdr_cells[4].text = 'Alerts'
        hdr_cells[5].text = 'Risk %'
        
        for classroom in classrooms:
            students = Student.objects.filter(enrollments__classroom=classroom, enrollments__is_active=True)
            total = students.count()
            high_risk = students.filter(risk_scores__risk_level__in=['high', 'critical']).distinct().count()
            alerts = Alert.objects.filter(student__in=students, status='active').count()
            risk_pct = (high_risk / total * 100) if total > 0 else 0
            
            row_cells = table.add_row().cells
            row_cells[0].text = classroom.name
            row_cells[1].text = classroom.form_master.name if classroom.form_master else 'N/A'
            row_cells[2].text = str(total)
            row_cells[3].text = str(high_risk)
            row_cells[4].text = str(alerts)
            row_cells[5].text = f'{risk_pct:.1f}%'
    
    def _generate_performance_content_docx(self, doc):
        """Generate performance content for DOCX"""
        doc.add_heading('Form Master Performance Metrics', 1)
        
        form_masters = User.objects.filter(role='form_master')
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Form Master'
        hdr_cells[1].text = 'Classroom'
        hdr_cells[2].text = 'Active Cases'
        hdr_cells[3].text = 'Avg Days'
        hdr_cells[4].text = 'On-Time %'
        hdr_cells[5].text = 'Rating'
        
        for fm in form_masters:
            classroom = Classroom.objects.filter(form_master=fm, is_active=True).first()
            cases = InterventionCase.objects.filter(assigned_to=fm)
            active = cases.exclude(status='closed').count()
            
            closed = cases.filter(status='closed')
            avg_resolution = closed.aggregate(avg_days=Avg(F('updated_at') - F('created_at')))['avg_days']
            avg_days = avg_resolution.days if avg_resolution else 0
            
            on_time = closed.filter(updated_at__lte=F('created_at') + timedelta(days=14)).count()
            on_time_pct = (on_time / closed.count() * 100) if closed.count() > 0 else 0
            
            rating = 'Excellent' if on_time_pct >= 80 else 'Good' if on_time_pct >= 60 else 'Fair'
            
            row_cells = table.add_row().cells
            row_cells[0].text = fm.name
            row_cells[1].text = classroom.name if classroom else 'N/A'
            row_cells[2].text = str(active)
            row_cells[3].text = str(round(avg_days, 1))
            row_cells[4].text = f'{on_time_pct:.1f}%'
            row_cells[5].text = rating
